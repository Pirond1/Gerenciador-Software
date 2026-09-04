"""
Geracao da ERS em .docx, seguindo o modelo da disciplina.

Word em vez de PDF de proposito: o professor consegue comentar e editar,
e o formato do modelo dele ja e .docx.

O documento e montado a partir dos mesmos JSONs que alimentam as telas --
nao existe conteudo que so exista no arquivo exportado.
"""

from __future__ import annotations

import re
from datetime import date
from io import BytesIO

from docx import Document as NovoDocumento
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.image.image import Image as ImagemDocx
from docx.shared import Cm, Pt

from app.models import PREFIXO_REQUISITO, TipoRequisito

# No modelo do professor estas tres sao irmas de "Funcoes do produto",
# em Heading 2 -- nao subsecoes numeradas dela.
PASTA_IMAGENS = None

SECOES_FUNCOES = [
    (TipoRequisito.BASICA, "FUNÇÕES BÁSICAS"),
    (TipoRequisito.FUNDAMENTAL, "FUNÇÕES FUNDAMENTAIS"),
    (TipoRequisito.SAIDA, "FUNÇÕES DE SAÍDA"),
]


# ---------------------------------------------------------------------------
# Auxiliares de formatacao
# ---------------------------------------------------------------------------


def _texto(doc, conteudo: str, *, italico: bool = False, negrito: bool = False):
    """Paragrafo com formatacao.

    `_texto(doc, x, italico=True)` nao faz nada -- Paragraph nao tem
    essa propriedade e o atributo e descartado sem erro. A formatacao vive
    no run, nao no paragrafo.
    """
    p = doc.add_paragraph()
    run = p.add_run(conteudo)
    run.italic = italico
    run.bold = negrito
    return p


def _numerar_titulos(doc) -> None:
    """Liga uma numeracao multinivel aos estilos Titulo 1/2/3.

    O modelo do professor nao escreve "1.1" no texto: o numero vem do
    estilo. Sem esta definicao, tirar os numeros manuais deixaria os
    titulos sem numeracao nenhuma.
    """
    numbering = doc.part.numbering_part.element

    abstrato = OxmlElement("w:abstractNum")
    abstrato.set(qn("w:abstractNumId"), "500")
    multinivel = OxmlElement("w:multiLevelType")
    multinivel.set(qn("w:val"), "multilevel")
    abstrato.append(multinivel)

    for nivel, formato in enumerate(["%1.", "%1.%2", "%1.%2.%3"]):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(nivel))
        for tag, valor in [("w:start", "1"), ("w:numFmt", "decimal"),
                           ("w:lvlText", formato), ("w:lvlJc", "left")]:
            el = OxmlElement(tag)
            el.set(qn("w:val"), valor)
            lvl.append(el)
        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:firstLine"), "0")
        ppr.append(ind)
        lvl.append(ppr)
        abstrato.append(lvl)

    numbering.insert(0, abstrato)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "500")
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), "500")
    num.append(ref)
    numbering.append(num)

    for nivel, nome in enumerate(["Heading 1", "Heading 2", "Heading 3"]):
        estilo = doc.styles[nome].element
        ppr = estilo.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            estilo.append(ppr)
        numpr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(nivel))
        numid = OxmlElement("w:numId")
        numid.set(qn("w:val"), "500")
        numpr.append(ilvl)
        numpr.append(numid)
        ppr.append(numpr)


def _sumario(doc) -> None:
    """Insere um campo de sumario.

    O Word so calcula o conteudo quando o usuario manda atualizar (Ctrl+A,
    F9). Nao da para preencher aqui: o numero de pagina so existe depois
    da paginacao, que quem faz e o Word.
    """
    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run()

    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instrucao = OxmlElement("w:instrText")
    instrucao.set(qn("xml:space"), "preserve")
    instrucao.text = 'TOC \\o "1-3" \\h \\z \\u'
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    aviso = OxmlElement("w:t")
    aviso.text = "Clique com o botão direito e escolha Atualizar campo."
    fim = OxmlElement("w:fldChar")
    fim.set(qn("w:fldCharType"), "end")

    for elemento in (inicio, instrucao, separador, aviso, fim):
        run._r.append(elemento)


def _markdown(doc, texto: str) -> None:
    """Converte o Markdown das telas em paragrafos do Word.

    Cobre o que as pessoas realmente usam ao escrever a ERS: paragrafos,
    listas com hifen e listas numeradas. Nao e um interpretador completo --
    tabela e bloco de codigo dentro de secao em prosa nao aparecem.
    """
    if not texto.strip():
        _texto(doc, "Seção ainda não redigida.", italico=True)
        return

    for bruto in texto.replace("\r\n", "\n").split("\n"):
        linha = bruto.strip()
        if not linha:
            continue
        if re.match(r"^[-*]\s+", linha):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", linha), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", linha):
            doc.add_paragraph(re.sub(r"^\d+[.)]\s+", "", linha), style="List Number")
        elif linha.startswith("#"):
            _texto(doc, linha.lstrip("# ").strip(), negrito=True)
        else:
            doc.add_paragraph(linha)


def _tabela(doc, cabecalho: list[str], linhas: list[list[str]], larguras: list[float]):
    """Tabela com largura fixa por coluna, em centimetros.

    Largura precisa ser definida na tabela E em cada celula, senao o Word
    redistribui as colunas sozinho e o layout muda de arquivo para arquivo.
    """
    tabela = doc.add_table(rows=1, cols=len(cabecalho))
    tabela.style = "Table Grid"
    tabela.autofit = False

    for i, largura in enumerate(larguras):
        tabela.columns[i].width = Cm(largura)

    for i, titulo in enumerate(cabecalho):
        celula = tabela.rows[0].cells[i]
        celula.text = ""
        run = celula.paragraphs[0].add_run(titulo)
        run.bold = True
        run.font.size = Pt(9)
        celula.width = Cm(larguras[i])

    for linha in linhas:
        celulas = tabela.add_row().cells
        for i, valor in enumerate(linha):
            celulas[i].text = ""
            run = celulas[i].paragraphs[0].add_run(str(valor))
            run.font.size = Pt(9)
            celulas[i].width = Cm(larguras[i])

    doc.add_paragraph()
    return tabela


def _figuras(doc, documento, secao: str) -> bool:
    """Insere as imagens de uma secao. Devolve True se inseriu alguma.

    A largura e limitada a 16 cm (a util da pagina A4 com as margens
    padrao) e a altura a 20 cm. Sem o limite de altura, um diagrama
    vertical estourado empurraria a legenda para a pagina seguinte.
    """
    figuras = documento.imagens.get(secao, [])
    if not figuras:
        return False

    for figura in figuras:
        caminho = PASTA_IMAGENS / figura.arquivo
        if not caminho.exists():
            _texto(doc, f"[imagem ausente: {figura.arquivo}]", italico=True)
            continue

        try:
            info = ImagemDocx.from_file(str(caminho))
            proporcao = info.px_height / info.px_width
            largura = min(Cm(16), Cm(20 / proporcao) if proporcao else Cm(16))
            doc.add_picture(str(caminho), width=largura)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            _texto(doc, f"[nao foi possivel inserir {figura.arquivo}: {e}]",
                   italico=True)
            continue

        if figura.legenda:
            legenda = doc.add_paragraph()
            legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = legenda.add_run(figura.legenda)
            run.italic = True
            run.font.size = Pt(9)

    return True


def _lista_rotulada(doc, rotulo: str, itens: list[str]) -> None:
    if not itens:
        return
    p = doc.add_paragraph()
    p.add_run(f"{rotulo}: ").bold = True
    for item in itens:
        doc.add_paragraph(item, style="List Bullet")


# ---------------------------------------------------------------------------
# Montagem do documento
# ---------------------------------------------------------------------------


def _capa(doc, projeto, equipe) -> None:
    for _ in range(4):
        doc.add_paragraph()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Especificação de Requisitos de Software")
    run.bold = True
    run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(projeto.nome)
    run.font.size = Pt(15)

    doc.add_paragraph()
    for texto in (projeto.disciplina,
                  f"Professor: {projeto.professor}" if projeto.professor else "",
                  projeto.semestre):
        if texto:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(texto)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Equipe").bold = True
    for membro in equipe.membros:
        if membro.perfil.value == "professor":
            continue
        papeis = ", ".join(x.value for x in membro.papeis)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{membro.nome}" + (f" — {papeis}" if papeis else ""))

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Gerado em {date.today():%d/%m/%Y}").italic = True

    doc.add_page_break()


def _requisito(doc, req) -> None:
    """Um paragrafo corrido, no formato do modelo:

    "RF_B01 – Gerenciar Cliente: O usuario podera... Itens obrigatorios: ..."

    Listas com marcador seriam mais legiveis na tela, mas o documento
    segue o modelo do professor, nao a nossa preferencia.
    """
    p = doc.add_paragraph()
    p.add_run(f"{req.id} – {req.titulo}: ").bold = True

    if req.descricao:
        p.add_run(re.sub(r"\s+", " ", req.descricao).strip() + " ")

    rotulo_entrada = ("Filtros de consulta" if req.tipo == TipoRequisito.SAIDA
                      else "Itens de informação obrigatórios")
    for rotulo, itens in [
        (rotulo_entrada, req.entradas),
        ("Itens de informação opcionais", req.opcionais),
        ("Itens de saída", req.saidas),
        ("Regras de validação", req.regras),
        ("Critérios de aceite", req.criterios_aceite),
    ]:
        if itens:
            p.add_run(f"{rotulo}: ").bold = True
            p.add_run("; ".join(itens) + ". ")

    p.add_run(f"(Prioridade: {req.prioridade.value}; situação: {req.status.value}"
              + (f"; origem: {req.origem}" if req.origem else "") + ".)").italic = True
    doc.add_paragraph()


def _caso_de_uso(doc, caso, atores_por_id) -> None:
    """Prosa com rotulos, como no modelo -- nao tabela.

    O modelo titula o caso pelo requisito que ele realiza; usamos o
    primeiro requisito referenciado e caimos no id do caso quando nao ha
    nenhum.
    """
    principal = caso.requisitos[0] if caso.requisitos else caso.id
    doc.add_heading(f"CASO DE USO: {principal} – {caso.nome}", level=3)

    ator = atores_por_id.get(caso.ator_principal)

    def rotulo(nome, valor=""):
        p = doc.add_paragraph()
        p.add_run(f"{nome}: ").bold = True
        if valor:
            p.add_run(valor)
        return p

    rotulo("Requisito", f"{principal} – {caso.nome}")
    rotulo("Ator Principal", ator.nome if ator else (caso.ator_principal or "—"))

    rotulo("Referências Cruzadas")
    if caso.requisitos:
        for req in caso.requisitos:
            doc.add_paragraph(f"- {req}")
    else:
        doc.add_paragraph("- Nenhuma")

    rotulo("Pré-condição")
    doc.add_paragraph(caso.pre_condicao or "—")
    rotulo("Pós-condição")
    doc.add_paragraph(caso.pos_condicao or "—")

    rotulo("Fluxo Principal")
    for i, passo in enumerate(caso.fluxo_principal, 1):
        prefixo = f"{passo.ator}: " if passo.ator else ""
        doc.add_paragraph(f"{i}. {prefixo}{passo.acao}")

    if caso.fluxos_alternativos:
        rotulo("Fluxo Alternativo")
        for fluxo in caso.fluxos_alternativos:
            _texto(doc, f"{fluxo.nome}:", negrito=True)
            for i, passo in enumerate(fluxo.passos, 1):
                prefixo = f"{passo.ator}: " if passo.ator else ""
                doc.add_paragraph(f"{i}. {prefixo}{passo.acao}")
    doc.add_paragraph()


def gerar(repo) -> BytesIO:
    global PASTA_IMAGENS
    PASTA_IMAGENS = repo.imagens_dir

    """Monta a ERS completa e devolve o arquivo em memoria."""
    projeto = repo.projeto()
    equipe = repo.equipe()
    documento = repo.documento()
    requisitos = repo.requisitos()
    atores = repo.atores().atores
    glossario = repo.glossario().termos
    casos = repo.casos_de_uso()
    atores_por_id = {a.id: a for a in atores}

    doc = NovoDocumento()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    _numerar_titulos(doc)

    _capa(doc, projeto, equipe)

    p = doc.add_paragraph()
    run = p.add_run("SUMÁRIO")
    run.bold = True
    run.font.size = Pt(14)
    _sumario(doc)
    doc.add_page_break()

    # ---------------- 1. Introducao ----------------
    doc.add_heading("INTRODUÇÃO", level=1)

    doc.add_heading("OBJETIVO", level=2)
    _markdown(doc, documento.objetivo)

    doc.add_heading("ESCOPO", level=2)
    _markdown(doc, documento.escopo)

    doc.add_heading("DEFINIÇÕES, SIGLAS E ABREVIAÇÕES", level=2)
    if glossario:
        _tabela(doc, ["Termo", "Definição"],
                [[t.termo, t.definicao] for t in glossario], [4.5, 11.5])
    else:
        _texto(doc, "Nenhum termo definido.", italico=True)

    doc.add_heading("REFERÊNCIAS", level=2)
    if documento.referencias:
        for ref in documento.referencias:
            doc.add_paragraph(ref, style="List Bullet")
    else:
        _texto(doc, "Nenhuma referência registrada.", italico=True)

    doc.add_heading("VISÃO GERAL", level=2)
    _markdown(doc, documento.visao_geral)
    doc.add_page_break()

    # ---------------- 2. Descricao geral ----------------
    doc.add_heading("DESCRIÇÃO GERAL DO PRODUTO", level=1)

    doc.add_heading("PERSPECTIVA DO PRODUTO", level=2)
    _markdown(doc, documento.perspectiva)

    doc.add_heading("FUNÇÕES DO PRODUTO", level=2)
    for tipo, titulo in SECOES_FUNCOES:
        doc.add_heading(titulo, level=2)
        do_tipo = [r for r in requisitos if r.tipo == tipo]
        if do_tipo:
            for req in do_tipo:
                _requisito(doc, req)
        else:
            _texto(doc, "Nenhuma função especificada.", italico=True)

    doc.add_heading("CARACTERÍSTICAS DO USUÁRIO", level=2)
    _markdown(doc, documento.caracteristicas_usuario)
    if atores:
        _tabela(
            doc, ["Ator", "Descrição", "Frequência", "Instrução", "Proficiência"],
            [[a.nome, a.descricao, a.frequencia_uso, a.nivel_instrucao, a.proficiencia]
             for a in atores],
            [3.2, 5.0, 2.6, 2.6, 2.6],
        )

    doc.add_heading("RESTRIÇÕES, DEPENDÊNCIAS E SUPOSIÇÕES", level=2)
    _markdown(doc, documento.restricoes)

    nao_funcionais = [r for r in requisitos if r.tipo == TipoRequisito.NAO_FUNCIONAL]
    if nao_funcionais:
        _texto(doc, "Requisitos não funcionais", negrito=True)
        for req in nao_funcionais:
            _requisito(doc, req)

    doc.add_heading("REQUISITOS ADIADOS", level=2)
    _markdown(doc, documento.requisitos_adiados)

    doc.add_heading("ESTUDO DE VIABILIDADE", level=2)
    _markdown(doc, documento.viabilidade)
    _figuras(doc, documento, "topologia")
    doc.add_page_break()

    # ---------------- 3. Requisitos especificos ----------------
    doc.add_heading("ANÁLISE ORIENTADA A OBJETOS", level=1)

    doc.add_heading("DIAGRAMAS DE CASO DE USO", level=2)
    if atores:
        doc.add_paragraph("Atores do sistema:")
        for a in atores:
            doc.add_paragraph(
                f"{a.nome}" + (f" — {a.descricao}" if a.descricao else ""),
                style="List Bullet")
    if not _figuras(doc, documento, "casos_uso"):
        _texto(doc, "Diagrama de casos de uso: inserir imagem.", italico=True)

    doc.add_heading("ESPECIFICAÇÕES DE CASO DE USO", level=2)
    if casos:
        for caso in casos:
            _caso_de_uso(doc, caso, atores_por_id)
    else:
        _texto(doc, "Nenhum caso de uso especificado.", italico=True)

    for chave, titulo in [("atividades", "DIAGRAMA DE ATIVIDADES"),
                          ("classes", "DIAGRAMA DE CLASSE"),
                          ("sequencia", "DIAGRAMA DE SEQUÊNCIA"),
                          ("der", "DIAGRAMA ENTIDADE RELACIONAMENTO"),
                          ("prototipo", "PROTÓTIPO")]:
        doc.add_heading(titulo, level=2)
        if not _figuras(doc, documento, chave):
            _texto(doc, "Inserir imagem.", italico=True)

    # ---------------- Apendice ----------------
    # A matriz nao esta no modelo do professor. Entra porque e o unico
    # lugar onde a rastreabilidade requisito -> tarefa fica visivel em
    # papel, e ela e o que a disciplina cobra.
    doc.add_page_break()
    doc.add_heading("APÊNDICES", level=1)
    _texto(doc, "Matriz de rastreabilidade", negrito=True)
    doc.add_paragraph(
        "Relação entre os requisitos especificados e as tarefas do projeto, "
        "extraída automaticamente do board.")

    if requisitos:
        linhas = []
        for req in requisitos:
            tarefas = repo.tarefas_do_requisito(req.id)
            linhas.append([
                req.id, req.titulo,
                ", ".join(t.id for t in tarefas) if tarefas else "SEM TAREFA",
            ])
        _tabela(doc, ["Requisito", "Título", "Tarefas"], linhas, [2.6, 7.4, 6.0])

    if documento.revisoes:
        doc.add_page_break()
        doc.add_heading("HISTÓRICO DE REVISÕES", level=1)
        _tabela(
            doc, ["Versão", "Data", "Responsável", "Histórico de Alterações"],
            [[r.versao, f"{r.data:%d/%m/%Y}", r.autor, r.descricao]
             for r in documento.revisoes],
            [2.2, 2.6, 3.4, 7.8],
        )

    memoria = BytesIO()
    doc.save(memoria)
    memoria.seek(0)
    return memoria